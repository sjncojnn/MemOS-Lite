"""Ingest tài liệu cho RAG baseline: đọc file -> chunk -> embed -> lưu vector.

Baseline này cố tình giữ đơn giản hơn MemOS-lite:
- Không MemoryUnit.
- Không lifecycle / TTL / hot-cold tier.
- Không provenance/version/conflict management.
- Chỉ lưu chunk text + metadata tối thiểu để đối chứng với nhánh MemOS-lite.
"""

from __future__ import annotations

import re
import unicodedata
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Iterable, Optional, Protocol, Sequence


class BaselineStoreLike(Protocol):
    def add(
        self,
        ids: list[str],
        embeddings: list[list[float]],
        documents: list[str],
        metadatas: list[dict[str, Any]],
    ) -> None: ...


@dataclass
class Chunk:
    """Đơn vị nhỏ nhất của RAG baseline.

    `metadata` chỉ chứa thông tin phẳng, tối thiểu để debug/eval nguồn truy hồi.
    Nó không phải provenance có vòng đời như MemOS-lite.
    """

    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    text: str = ""
    source: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_WHITESPACE_RE = re.compile(r"[ \t\r\f\v]+")
_MULTI_NEWLINE_RE = re.compile(r"\n{3,}")


def clean_text(raw_text: Any) -> str:
    """Chuẩn hóa text nhẹ để baseline không vỡ vì ký tự lỗi/space thừa."""

    if raw_text is None:
        return ""
    text = str(raw_text)

    # Optional dependency; nếu không có thì bỏ qua để baseline vẫn chạy được.
    try:
        import ftfy  # type: ignore

        text = ftfy.fix_text(text)
    except Exception:
        pass

    text = unicodedata.normalize("NFC", text.replace("\xa0", " "))
    text = _CONTROL_CHARS_RE.sub(" ", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(_WHITESPACE_RE.sub(" ", line).strip() for line in text.split("\n"))
    text = _MULTI_NEWLINE_RE.sub("\n\n", text)
    return text.strip()


def _norm_header(value: Any) -> str:
    text = clean_text(value).lower().replace("đ", "d")
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _find_col(headers: list[str], candidates: Iterable[str]) -> Optional[int]:
    normalized_candidates = [_norm_header(c) for c in candidates]
    for idx, header in enumerate(headers):
        if header in normalized_candidates:
            return idx
    for idx, header in enumerate(headers):
        if any(candidate in header for candidate in normalized_candidates):
            return idx
    return None


def _category_columns(headers: list[str], question_idx: int, answer_idx: int) -> list[int]:
    keys = (
        "category",
        "topic",
        "nhom",
        "phan loai",
        "chu de",
        "cap",
        "level",
        "loai",
        "nghiep vu",
        "kenh",
        "san pham",
    )
    return [
        idx
        for idx, header in enumerate(headers)
        if idx not in {question_idx, answer_idx} and any(key in header for key in keys)
    ]


def _stable_chunk_id(source: str, index: int, text: str) -> str:
    """ID ổn định hơn UUID để upsert baseline không phình khi ingest lại cùng dữ liệu."""

    import hashlib

    digest = hashlib.sha1(f"{source}\n{index}\n{text}".encode("utf-8")).hexdigest()[:16]
    return f"baseline-{digest}"


def chunk_text(text: str, chunk_size_chars: int = 1000, overlap_chars: int = 100) -> list[str]:
    """Chunk theo ký tự với overlap cố định.

    Đây là chunking đơn giản cho RAG baseline, không semantic chunking theo heading như
    MemOS-lite. Hàm ưu tiên cắt theo ranh giới đoạn gần cuối cửa sổ để chunk dễ đọc hơn,
    nhưng vẫn fallback sang cắt ký tự khi đoạn quá dài.
    """

    text = clean_text(text)
    if not text:
        return []
    if chunk_size_chars <= 0:
        raise ValueError("chunk_size_chars phải > 0")
    if overlap_chars < 0:
        raise ValueError("overlap_chars phải >= 0")
    if chunk_size_chars <= overlap_chars:
        raise ValueError("chunk_size_chars phải lớn hơn overlap_chars")
    if len(text) <= chunk_size_chars:
        return [text]

    chunks: list[str] = []
    start = 0
    step = chunk_size_chars - overlap_chars

    while start < len(text):
        hard_end = min(start + chunk_size_chars, len(text))
        end = hard_end

        # Nếu chưa tới cuối văn bản, cố gắng cắt ở cuối đoạn/câu gần hard_end.
        if hard_end < len(text):
            window = text[start:hard_end]
            paragraph_cut = max(window.rfind("\n\n"), window.rfind("\n"))
            sentence_cut = max(window.rfind(". "), window.rfind("? "), window.rfind("! "))
            cut = max(paragraph_cut, sentence_cut)
            # Chỉ dùng điểm cắt nếu nó không làm chunk quá ngắn.
            if cut >= int(chunk_size_chars * 0.55):
                end = start + cut + (2 if cut == sentence_cut else 0)

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break
        start = max(end - overlap_chars, start + step)

    return chunks


def _docx_table_text(table: Any) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        cells = [cell for cell in cells if cell]
        if cells:
            rows.append(" | ".join(cells))
    return clean_text("\n".join(rows))


def load_docx_as_chunks(
    file_path: str | Path,
    *,
    chunk_size_chars: int = 1000,
    overlap_chars: int = 100,
) -> list[Chunk]:
    """Đọc 1 file .docx thành các chunk text tối giản.

    Paragraph và table đều được gom vào cùng text rồi chunk theo `chunk_text()`.
    """

    from docx import Document

    file_path = Path(file_path)
    doc = Document(file_path)

    parts: list[str] = []
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if text:
            parts.append(text)

    for table in doc.tables:
        table_text = _docx_table_text(table)
        if table_text:
            parts.append(table_text)

    full_text = clean_text("\n\n".join(parts))
    source = file_path.name
    chunks: list[Chunk] = []
    for idx, text in enumerate(chunk_text(full_text, chunk_size_chars, overlap_chars), start=1):
        chunks.append(
            Chunk(
                id=_stable_chunk_id(source, idx, text),
                text=text,
                source=source,
                metadata={"source": source, "file_name": file_path.name, "chunk_index": idx, "file_type": "docx"},
            )
        )
    return chunks


def load_faq_xlsx_as_chunks(file_path: str | Path) -> list[Chunk]:
    """Đọc FAQ .xlsx, mỗi dòng Q&A là 1 chunk.

    Tự dò cột câu hỏi/trả lời theo header phổ biến. Nếu không dò được, fallback sang
    hai cột đầu tiên.
    """

    from openpyxl import load_workbook

    file_path = Path(file_path)
    wb = load_workbook(file_path, read_only=True, data_only=True)

    question_names = [
        "question",
        "q",
        "cau hoi",
        "câu hỏi",
        "noi dung cau hoi",
        "nội dung câu hỏi",
        "cau hoi thuong gap",
        "câu hỏi thường gặp",
    ]
    answer_names = [
        "answer",
        "a",
        "tra loi",
        "trả lời",
        "noi dung tra loi",
        "nội dung trả lời",
        "dap an",
        "đáp án",
    ]

    chunks: list[Chunk] = []
    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        headers = [_norm_header(cell) for cell in rows[0]]
        q_idx = _find_col(headers, question_names)
        a_idx = _find_col(headers, answer_names)
        if q_idx is None or a_idx is None:
            q_idx, a_idx = 0, 1

        cat_cols = _category_columns(headers, q_idx, a_idx)

        for excel_row_idx, row in enumerate(rows[1:], start=2):
            row_values = list(row)
            question = clean_text(row_values[q_idx] if q_idx < len(row_values) else "")
            answer = clean_text(row_values[a_idx] if a_idx < len(row_values) else "")
            if not question and not answer:
                continue

            categories: list[str] = []
            for col in cat_cols:
                value = clean_text(row_values[col] if col < len(row_values) else "")
                if value:
                    categories.append(value)

            category_text = " > ".join(categories)
            text_parts = []
            if category_text:
                text_parts.append(f"Phân loại: {category_text}")
            if question:
                text_parts.append(f"Câu hỏi: {question}")
            if answer:
                text_parts.append(f"Trả lời: {answer}")
            text = clean_text("\n".join(text_parts))

            source = f"{file_path.name}::{sheet.title}::row:{excel_row_idx}"
            chunks.append(
                Chunk(
                    id=_stable_chunk_id(source, 1, text),
                    text=text,
                    source=source,
                    metadata={
                        "source": source,
                        "file_name": file_path.name,
                        "sheet": sheet.title,
                        "row": excel_row_idx,
                        "file_type": "xlsx",
                        "category": category_text,
                        "question": question,
                    },
                )
            )

    return chunks


def load_path_as_chunks(
    path: str | Path,
    *,
    chunk_size_chars: int = 1000,
    overlap_chars: int = 100,
) -> list[Chunk]:
    """Load một file DOCX/XLSX thành chunks."""

    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return load_docx_as_chunks(path, chunk_size_chars=chunk_size_chars, overlap_chars=overlap_chars)
    if suffix in {".xlsx", ".xlsm"}:
        return load_faq_xlsx_as_chunks(path)
    raise ValueError(f"Không hỗ trợ file baseline: {path}")


def load_paths_as_chunks(
    paths: Sequence[str | Path],
    *,
    chunk_size_chars: int = 1000,
    overlap_chars: int = 100,
) -> list[Chunk]:
    """Load nhiều file/folder thành chunks.

    Nếu truyền folder, hàm đọc đệ quy *.docx, *.xlsx, *.xlsm.
    """

    all_files: list[Path] = []
    for raw_path in paths:
        path = Path(raw_path)
        if path.is_dir():
            all_files.extend(sorted(path.rglob("*.docx")))
            all_files.extend(sorted(path.rglob("*.xlsx")))
            all_files.extend(sorted(path.rglob("*.xlsm")))
        elif path.is_file():
            all_files.append(path)
        else:
            raise FileNotFoundError(f"Không tìm thấy path: {path}")

    chunks: list[Chunk] = []
    for file_path in all_files:
        chunks.extend(
            load_path_as_chunks(
                file_path,
                chunk_size_chars=chunk_size_chars,
                overlap_chars=overlap_chars,
            )
        )
    return chunks


def _embed_many(texts: list[str], embed_fn: Callable[[str], list[float]] | Any) -> list[list[float]]:
    """Dùng batch embed nếu client có embed_many(), nếu không gọi từng text."""

    embed_many = getattr(embed_fn, "embed_many", None)
    if callable(embed_many):
        return [[float(x) for x in emb] for emb in embed_many(texts)]
    return [[float(x) for x in embed_fn(text)] for text in texts]


def ingest_and_embed(
    chunks: list[Chunk],
    embed_fn: Callable[[str], list[float]],
    store: BaselineStoreLike,
    *,
    batch_size: int = 64,
) -> dict[str, int]:
    """Embed chunks và ghi vào baseline vector store.

    Returns thống kê đơn giản để script/CLI dễ in kết quả.
    """

    clean_chunks = [c for c in chunks if clean_text(c.text)]
    if not clean_chunks:
        return {"chunks": 0, "inserted": 0}
    if batch_size <= 0:
        raise ValueError("batch_size phải > 0")

    inserted = 0
    for start in range(0, len(clean_chunks), batch_size):
        batch = clean_chunks[start : start + batch_size]
        ids = [c.id for c in batch]
        texts = [clean_text(c.text) for c in batch]
        embeddings = _embed_many(texts, embed_fn)
        metadatas = []
        for c in batch:
            metadata = dict(c.metadata)
            metadata.setdefault("source", c.source)
            metadatas.append(metadata)
        store.add(ids=ids, embeddings=embeddings, documents=texts, metadatas=metadatas)
        inserted += len(batch)

    return {"chunks": len(clean_chunks), "inserted": inserted}
