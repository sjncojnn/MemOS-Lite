"""Load business DOCX files and FAQ XLSX files into raw MemoryUnit objects.

The module stays intentionally simple: DOCX is split by headings/sections and
FAQ is split one row = one MemoryUnit. Storing, dedup, conflict detection, and
vector indexing are handled by memory_store.py and memory_ops.py.
"""

from __future__ import annotations

import re
import unicodedata
from datetime import timedelta
from pathlib import Path
from typing import Any, Iterable, Optional

from src.config import Config
from src.schemas import MemoryTier, MemoryUnit, Provenance, SourceType, now_utc

CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
WHITESPACE_RE = re.compile(r"[ \t\r\f\v]+")
MULTI_NEWLINE_RE = re.compile(r"\n{3,}")


def clean_text(raw_text: Any) -> str:
    """Normalize whitespace and remove common invisible/control characters."""

    if raw_text is None:
        return ""
    text = str(raw_text)

    # Use ftfy when available, but keep it optional for a small dependency set.
    try:
        import ftfy  # type: ignore

        text = ftfy.fix_text(text)
    except Exception:
        pass

    text = unicodedata.normalize("NFC", text.replace("\xa0", " "))
    text = CONTROL_CHARS_RE.sub(" ", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(WHITESPACE_RE.sub(" ", line).strip() for line in text.split("\n"))
    text = MULTI_NEWLINE_RE.sub("\n\n", text)
    return text.strip()


def build_provenance(source_type: SourceType, source_path: str | Path, source_ref: str = "") -> Provenance:
    """Create minimal provenance for one MemoryUnit."""

    return Provenance(
        source_type=source_type,
        source_path=str(source_path),
        source_ref=source_ref,
    )


def _ttl_expires_at(ttl_seconds: Optional[int]):
    if ttl_seconds is None or ttl_seconds <= 0:
        return None
    return now_utc() + timedelta(seconds=ttl_seconds)


def _slug_header(text: str) -> str:
    return clean_text(text).strip(" #:-–—") or "section"


def _rough_max_chars(chunk_size_tokens: Optional[int]) -> int:
    # Rough but practical for Vietnamese/English text; avoids tokenizer dependency.
    return max(800, int((chunk_size_tokens or 1200) * 4))


def _rough_overlap_chars(chunk_overlap_tokens: Optional[int]) -> int:
    return max(0, int((chunk_overlap_tokens or 0) * 4))


def _split_long_text(text: str, max_chars: int, overlap_chars: int = 0) -> list[str]:
    """Split only when a section is too long; prefer paragraph boundaries."""

    text = clean_text(text)
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]

    paragraphs = [p for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        if not current:
            current = para
            continue
        if len(current) + len(para) + 2 <= max_chars:
            current += "\n\n" + para
        else:
            chunks.append(current.strip())
            if overlap_chars > 0:
                current = current[-overlap_chars:].strip() + "\n\n" + para
            else:
                current = para

    if current.strip():
        chunks.append(current.strip())

    # Fallback for a single huge paragraph.
    final_chunks: list[str] = []
    for chunk in chunks:
        if len(chunk) <= max_chars:
            final_chunks.append(chunk)
            continue
        start = 0
        step = max(1, max_chars - overlap_chars)
        while start < len(chunk):
            final_chunks.append(chunk[start : start + max_chars].strip())
            start += step
    return [c for c in final_chunks if c]


def _make_unit(
    *,
    content: str,
    source: str,
    source_type: SourceType,
    source_path: str | Path,
    source_ref: str,
    category: str,
    tags: Optional[list[str]] = None,
    tier: MemoryTier = MemoryTier.WARM,
    ttl_seconds: Optional[int] = None,
    extra_metadata: Optional[dict[str, Any]] = None,
) -> MemoryUnit:
    return MemoryUnit(
        content=clean_text(content),
        source=source,
        category=category or "uncategorized",
        tags=tags or [],
        tier=tier,
        ttl_expires_at=_ttl_expires_at(ttl_seconds),
        provenance=build_provenance(source_type, source_path, source_ref),
        extra_metadata=extra_metadata or {},
    )


def load_docx_as_units(
    file_path: str | Path,
    category: str = "uncategorized",
    *,
    chunk_size_tokens: int = 1200,
    chunk_overlap_tokens: int = 150,
    ttl_seconds: Optional[int] = None,
) -> list[MemoryUnit]:
    """Read one .docx business document into section-level MemoryUnit objects."""

    from docx import Document

    file_path = Path(file_path)
    doc = Document(file_path)
    max_chars = _rough_max_chars(chunk_size_tokens)
    overlap_chars = _rough_overlap_chars(chunk_overlap_tokens)

    units: list[MemoryUnit] = []
    current_heading = file_path.stem
    buffer: list[str] = []
    section_index = 0

    def flush() -> None:
        nonlocal buffer, section_index
        body = clean_text("\n\n".join(buffer))
        buffer = []
        if not body:
            return
        section_index += 1
        heading = _slug_header(current_heading)
        content = f"{heading}\n\n{body}" if heading else body
        for part_idx, chunk in enumerate(_split_long_text(content, max_chars, overlap_chars), start=1):
            ref = f"section:{section_index}:{heading}"
            if part_idx > 1:
                ref += f":part:{part_idx}"
            units.append(
                _make_unit(
                    content=chunk,
                    source=file_path.name,
                    source_type=SourceType.DOC,
                    source_path=file_path,
                    source_ref=ref,
                    category=category,
                    tags=[category, file_path.stem, heading],
                    tier=MemoryTier.WARM,
                    ttl_seconds=ttl_seconds,
                    extra_metadata={"doc_heading": heading, "chunk_part": part_idx},
                )
            )

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue
        style_norm = _norm_header(para.style.name or "")
        is_heading = (
            style_norm.startswith("heading")
            or style_norm.startswith("title")
            or style_norm.startswith("tieu de")
        )
        if is_heading:
            flush()
            current_heading = text
        else:
            buffer.append(text)
    flush()

    # Tables are usually procedural/fee data; keep each table as one compact unit.
    for table_idx, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        table_text = clean_text("\n".join(rows))
        if not table_text:
            continue
        for part_idx, chunk in enumerate(_split_long_text(table_text, max_chars, overlap_chars), start=1):
            units.append(
                _make_unit(
                    content=chunk,
                    source=file_path.name,
                    source_type=SourceType.DOC,
                    source_path=file_path,
                    source_ref=f"table:{table_idx}:part:{part_idx}",
                    category=category,
                    tags=[category, file_path.stem, "table"],
                    tier=MemoryTier.WARM,
                    ttl_seconds=ttl_seconds,
                    extra_metadata={"doc_table": table_idx, "chunk_part": part_idx},
                )
            )

    return units


def _norm_header(value: Any) -> str:
    text = clean_text(value).lower()
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _find_col(headers: list[str], candidates: Iterable[str]) -> Optional[int]:
    normalized_candidates = [_norm_header(c) for c in candidates]
    for idx, header in enumerate(headers):
        if header in normalized_candidates:
            return idx
    for idx, header in enumerate(headers):
        if any(candidate in header for candidate in normalized_candidates):
            return idx
    return None


def _category_columns(headers: list[str], question_idx: int, answer_idx: int) -> list[int]:
    keys = ("category", "topic", "nhom", "phan loai", "chu de", "cap", "level", "loai", "nghiep vu", "kenh", "san pham",)
    cols = [
        idx
        for idx, header in enumerate(headers)
        if idx not in {question_idx, answer_idx} and any(key in header for key in keys)
    ]
    return cols


def load_faq_xlsx_as_units(
    file_path: str | Path,
    category: str = "faq",
    *,
    ttl_seconds: Optional[int] = None,
) -> list[MemoryUnit]:
    """Read FAQ .xlsx into MemoryUnit objects, one row per Q/A pair."""

    from openpyxl import load_workbook

    file_path = Path(file_path)
    wb = load_workbook(file_path, read_only=True, data_only=True)
    units: list[MemoryUnit] = []

    question_names = ["question", "q", "cau hoi", "câu hỏi", "noi dung cau hoi", "nội dung câu hỏi", "cau hoi thuong gap", "câu hỏi thường gặp"]
    answer_names = ["answer", "a", "tra loi", "trả lời", "noi dung tra loi", "nội dung trả lời", "dap an", "đáp án"]

    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        headers = [_norm_header(cell) for cell in rows[0]]
        q_idx = _find_col(headers, question_names)
        a_idx = _find_col(headers, answer_names)

        if q_idx is None or a_idx is None:
            # Common fallback: first two non-empty columns are question/answer.
            q_idx, a_idx = 0, 1

        cat_cols = _category_columns(headers, q_idx, a_idx)

        for excel_row_idx, row in enumerate(rows[1:], start=2):
            row_values = list(row)
            question = clean_text(row_values[q_idx] if q_idx < len(row_values) else "")
            answer = clean_text(row_values[a_idx] if a_idx < len(row_values) else "")
            if not question or not answer:
                continue

            category_parts = [
                clean_text(row_values[idx])
                for idx in cat_cols
                if idx < len(row_values) and clean_text(row_values[idx])
            ]
            unit_category = "/".join(category_parts) if category_parts else category
            tags = ["faq", *category_parts]
            content = f"Q: {question}\nA: {answer}"

            units.append(
                _make_unit(
                    content=content,
                    source=file_path.name,
                    source_type=SourceType.FAQ,
                    source_path=file_path,
                    source_ref=f"sheet:{sheet.title}:row:{excel_row_idx}",
                    category=unit_category,
                    tags=tags,
                    tier=MemoryTier.HOT,
                    ttl_seconds=ttl_seconds,
                    extra_metadata={
                        "question": question,
                        "answer": answer,
                        "sheet": sheet.title,
                        "row": excel_row_idx,
                    },
                )
            )

    return units


def load_file_as_units(file_path: str | Path, config: Optional[Config] = None, category: str = "uncategorized") -> list[MemoryUnit]:
    """Dispatch loader by suffix."""

    config = config or Config()
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return load_docx_as_units(
            file_path,
            category=category,
            chunk_size_tokens=config.chunk_size_tokens,
            chunk_overlap_tokens=config.chunk_overlap_tokens,
            ttl_seconds=config.default_ttl_seconds,
        )
    if suffix in {".xlsx", ".xlsm"}:
        return load_faq_xlsx_as_units(
            file_path,
            category="faq" if category == "uncategorized" else category,
            ttl_seconds=config.default_ttl_seconds,
        )
    raise ValueError(f"Unsupported ingest file type: {file_path.suffix}")


def load_path_as_units(path: str | Path, config: Optional[Config] = None, category: str = "uncategorized") -> list[MemoryUnit]:
    """Load one supported file or all supported files in a directory."""

    path = Path(path)
    if path.is_file():
        return load_file_as_units(path, config=config, category=category)

    units: list[MemoryUnit] = []
    for file_path in sorted(path.rglob("*")):
        if file_path.suffix.lower() in {".docx", ".xlsx", ".xlsm"} and not file_path.name.startswith("~$"):
            units.extend(load_file_as_units(file_path, config=config, category=category))
    return units
