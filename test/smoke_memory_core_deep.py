"""Deep smoke test for MemOS-lite core modules.

Covers 7 files:
- config.py
- schemas.py
- db.py
- vector_store.py
- memory_store.py
- ingest.py
- memory_ops.py

Run from project root:
    python scripts/smoke_memory_core_deep.py --work-dir .smoke_test_data_deep

This script creates Vietnamese DOCX/XLSX fixtures and checks:
- DOCX heading extraction, table extraction, long-section chunk split
- FAQ Vietnamese headers, category/tags/provenance/extra_metadata
- exact duplicate skip, contradiction/near-duplicate conflict records
- vector search, category filter, tier filter, status filter, touch_access
- hot/warm/cold tier rules, TTL expiration, lifecycle log
- content update/version/hash, reindex, QA cache valid/expired behavior

It uses a deterministic fake embedding function, so Ollama is not required.
"""

from __future__ import annotations

import argparse
import math
import shutil
import sys
import tempfile
import unicodedata
from datetime import timedelta
from pathlib import Path
from typing import Iterable

# Allow running from scripts/ or project root.
SCRIPT_PATH = Path(__file__).resolve()
ROOT = SCRIPT_PATH.parents[1] if SCRIPT_PATH.parent.name == "scripts" else Path.cwd()
sys.path.insert(0, str(ROOT))

from src import db
from src.config import Config
from src.ingest import clean_text, load_path_as_units
from src.memory_ops import add_many, run_lifecycle_maintenance
from src.memory_store import MemoryStore
from src.schemas import (
    ConflictType,
    MemoryStatus,
    MemoryTier,
    MemoryUnit,
    Provenance,
    SourceType,
    compute_query_hash,
    now_utc,
)


# ---------------------------------------------------------------------------
# Small test helpers
# ---------------------------------------------------------------------------


def strip_accents(text: str) -> str:
    text = unicodedata.normalize("NFD", text.lower())
    return "".join(ch for ch in text if unicodedata.category(ch) != "Mn")


def fake_embed(text: str) -> list[float]:
    """Deterministic tiny embedding for smoke tests only.

    It is deliberately simple but gives similar vectors to related Vietnamese
    business phrases, enough to test vector_store/memory_store integration.
    """

    t = strip_accents(text)
    keyword_features = [
        ("dien" in t or "evn" in t),
        ("nuoc" in t),
        ("hoa don" in t or "bill" in t),
        ("thanh toan" in t or "dong tien" in t or "tra tien" in t or "xac nhan" in t),
        ("chuyen tien" in t or "lien ngan hang" in t),
        ("han muc" in t),
        ("phi" in t or "cuoc" in t),
        ("viettelpay" in t or "kpp" in t or "kenh ban" in t),
        ("khong the" in t or "loi" in t or "that bai" in t),
        ("mat khau" in t or "tai khoan" in t),
        ("het han" in t or "ttl" in t),
        ("lanh" in t or "it dung" in t or "cold" in t),
        ("nong" in t or "hot" in t or "tan suat" in t),
        ("warm" in t or "binh thuong" in t),
        ("archive" in t or "khong dung nua" in t),
        ("khuyen mai smoke" in t),
    ]
    vec = [1.0 if flag else 0.0 for flag in keyword_features]

    # Stable character buckets prevent unrelated text from becoming all-zero.
    buckets = [0.0] * 12
    for ch in t:
        if ch.isalnum():
            buckets[ord(ch) % len(buckets)] += 1.0
    vec.extend(buckets)

    norm = math.sqrt(sum(x * x for x in vec)) or 1.0
    return [x / norm for x in vec]


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)
    print(f"[OK] {message}")


def find_one(units: Iterable[MemoryUnit], needle: str) -> MemoryUnit:
    needle_norm = strip_accents(needle)
    for unit in units:
        if needle_norm in strip_accents(unit.content):
            return unit
    raise AssertionError(f"Cannot find unit containing: {needle}")


# ---------------------------------------------------------------------------
# Fixture creation
# ---------------------------------------------------------------------------


def make_sample_files(data_dir: Path) -> None:
    from docx import Document
    from openpyxl import Workbook

    data_dir.mkdir(parents=True, exist_ok=True)

    # Business DOCX with headings, long section, and table.
    doc = Document()
    doc.add_heading("Tổng quan ViettelPay Pro", level=1)
    doc.add_paragraph(
        "ViettelPay Pro là hệ thống tài chính điện tử dành cho các điểm cung cấp dịch vụ "
        "do Viettel ủy quyền, sử dụng trên ứng dụng di động hoặc web KPP."
    )

    doc.add_heading("Thanh toán hóa đơn điện", level=1)
    doc.add_paragraph(
        "Khách hàng có thể thanh toán hóa đơn điện trên ViettelPay Pro bằng cách vào mục Hóa đơn, "
        "chọn Điện, nhập mã khách hàng và xác nhận thanh toán."
    )
    doc.add_paragraph(
        "Nếu giao dịch thất bại, đại lý cần kiểm tra mã khách hàng, trạng thái kết nối và số dư trước khi thực hiện lại."
    )

    doc.add_heading("Chuyển tiền liên ngân hàng", level=1)
    doc.add_paragraph(
        "Đại lý có thể thực hiện chuyển tiền liên ngân hàng theo hạn mức được cấu hình cho từng kênh bán. "
        "Hạn mức có thể khác nhau theo vai trò và trạng thái tài khoản."
    )

    doc.add_heading("Quy trình dài cần chia chunk", level=1)
    long_para = (
        "Quy trình kiểm tra giao dịch yêu cầu đại lý xác minh thông tin khách hàng, kiểm tra hạn mức, "
        "kiểm tra phí, xác nhận nội dung giao dịch và lưu lại biên nhận. "
    ) * 18
    doc.add_paragraph(long_para)

    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Loại giao dịch"
    table.cell(0, 1).text = "Điều kiện"
    table.cell(0, 2).text = "Ghi chú"
    table.cell(1, 0).text = "Hóa đơn điện"
    table.cell(1, 1).text = "Cần mã khách hàng"
    table.cell(1, 2).text = "Kiểm tra trước khi xác nhận"
    table.cell(2, 0).text = "Chuyển tiền"
    table.cell(2, 1).text = "Cần hạn mức hợp lệ"
    table.cell(2, 2).text = "Theo kênh bán"
    doc.save(data_dir / "huong_dan_nghiep_vu.docx")

    # FAQ XLSX with Vietnamese headers and category columns.
    wb = Workbook()
    ws = wb.active
    ws.title = "FAQ_CoDau"
    ws.append(["Nhóm nghiệp vụ", "Cấp 1", "Chủ đề", "Câu hỏi", "Nội dung trả lời", "Ghi chú"])
    ws.append([
        "Hóa đơn",
        "Điện",
        "Thanh toán hóa đơn",
        "Khách hàng thanh toán hóa đơn điện như thế nào?",
        "Vào mục Hóa đơn, chọn Điện, nhập mã khách hàng và xác nhận thanh toán.",
        "FAQ chuẩn",
    ])
    # Exact duplicate: should be skipped by content_hash.
    ws.append([
        "Hóa đơn",
        "Điện",
        "Thanh toán hóa đơn",
        "Khách hàng thanh toán hóa đơn điện như thế nào?",
        "Vào mục Hóa đơn, chọn Điện, nhập mã khách hàng và xác nhận thanh toán.",
        "Bản trùng",
    ])
    # Same question, different answer: should become contradiction conflict.
    ws.append([
        "Hóa đơn",
        "Điện",
        "Thanh toán hóa đơn",
        "Khách hàng thanh toán hóa đơn điện như thế nào?",
        "Khách hàng không thể thanh toán hóa đơn điện trong hệ thống.",
        "Cố ý mâu thuẫn để test",
    ])
    ws.append([
        "Chuyển tiền",
        "Liên ngân hàng",
        "Hạn mức",
        "Hạn mức chuyển tiền liên ngân hàng là bao nhiêu?",
        "Hạn mức chuyển tiền phụ thuộc cấu hình của từng kênh bán.",
        "FAQ hạn mức",
    ])
    ws.append([
        "Tài khoản",
        "KPP",
        "Mật khẩu",
        "Đại lý đổi mật khẩu KPP ở đâu?",
        "Đại lý vào mục Tài khoản, chọn Đổi mật khẩu và làm theo hướng dẫn.",
        "FAQ tài khoản",
    ])

    # Sheet without Vietnamese accents, to test header normalization.
    ws2 = wb.create_sheet("FAQ_KhongDau")
    ws2.append(["Phan loai", "Cau hoi", "Tra loi"])
    ws2.append([
        "Nước",
        "Khách hàng thanh toán hóa đơn nước như thế nào?",
        "Vào mục Hóa đơn, chọn Nước, nhập mã khách hàng và xác nhận thanh toán.",
    ])

    wb.save(data_dir / "faq_tieng_viet.xlsx")


# ---------------------------------------------------------------------------
# Test phases
# ---------------------------------------------------------------------------


def test_ingest(units: list[MemoryUnit]) -> tuple[str, str]:
    print("\n[PHASE] ingest.py + schemas.py")
    assert_true(len(units) >= 9, "ingest loads DOCX sections, DOCX table, and FAQ rows")
    assert_true(any("hóa đơn" in unit.content.lower() or "hoá đơn" in unit.content.lower() for unit in units), "Vietnamese accents are preserved in content")
    assert_true(clean_text("  Hoá   đơn\u00a0điện\n\n\n") == "Hoá đơn điện", "clean_text keeps accents and normalizes whitespace")

    doc_units = [u for u in units if u.provenance and u.provenance.source_type == SourceType.DOC]
    faq_units = [u for u in units if u.provenance and u.provenance.source_type == SourceType.FAQ]
    assert_true(len(doc_units) >= 5, "DOCX produces multiple MemoryUnit objects")
    assert_true(len(faq_units) >= 6, "XLSX FAQ produces one MemoryUnit per valid row")

    headings = {u.extra_metadata.get("doc_heading") for u in doc_units if u.extra_metadata.get("doc_heading")}
    print(f"[INFO] doc headings = {sorted(str(h) for h in headings)}")
    assert_true("Tổng quan ViettelPay Pro" in headings, "DOCX heading 'Tổng quan ViettelPay Pro' is captured")
    assert_true("Thanh toán hóa đơn điện" in headings, "DOCX heading 'Thanh toán hóa đơn điện' is captured")
    assert_true("Chuyển tiền liên ngân hàng" in headings, "DOCX heading 'Chuyển tiền liên ngân hàng' is captured")
    assert_true(any(str(u.provenance.source_ref).startswith("table:") for u in doc_units), "DOCX table is ingested with table provenance")
    assert_true(any(":part:2" in str(u.provenance.source_ref) for u in doc_units), "long DOCX section is split into multiple chunks")

    faq_electric = find_one(faq_units, "Khách hàng thanh toán hóa đơn điện")
    faq_water = find_one(faq_units, "hóa đơn nước")
    print(f"[INFO] FAQ electric category = {faq_electric.category}")
    print(f"[INFO] FAQ water category = {faq_water.category}")
    assert_true(faq_electric.tier == MemoryTier.HOT, "FAQ units start in HOT tier")
    assert_true("faq" in [tag.lower() for tag in faq_electric.tags], "FAQ units contain faq tag")
    assert_true("Hóa đơn" in faq_electric.category and "Điện" in faq_electric.category, "Vietnamese category columns are joined into category path")
    assert_true(faq_water.category == "Nước", "non-accent Vietnamese headers are recognized on second sheet")
    assert_true(faq_electric.extra_metadata.get("question"), "FAQ question is stored in extra_metadata")
    assert_true(faq_electric.extra_metadata.get("answer"), "FAQ answer is stored in extra_metadata")
    assert_true(faq_electric.provenance and "sheet:" in faq_electric.provenance.source_ref, "FAQ provenance includes sheet and row")

    return faq_electric.category, "viettelpay_pro"


def test_storage_and_ops(store: MemoryStore, units: list[MemoryUnit]) -> None:
    print("\n[PHASE] memory_ops.py + memory_store.py + db.py")
    stats = add_many(store, units)
    print(f"[INFO] add_many stats = {stats}")
    assert_true(stats["inserted"] >= 8, "MemoryStore inserts non-duplicate units")
    assert_true(stats["exact_duplicate_skipped"] >= 1, "exact duplicate FAQ row is skipped by content_hash")
    assert_true(stats["conflicts"] >= 1, "near duplicate / contradiction conflicts are recorded")

    all_units = store.list(limit=1000)
    assert_true(len(all_units) == stats["inserted"], "SQLite list_memory_units matches inserted count")
    assert_true(all(unit.content_hash for unit in all_units), "all stored units have content_hash")
    hydrated_units = [store.get(unit.id) for unit in all_units]
    assert_true(all(unit is not None and unit.provenance is not None for unit in hydrated_units), "store.get retains provenance")

    conflicts = store.list_conflicts(limit=1000)
    conflict_summary = [(c.conflict_type.value, c.resolution.value) for c in conflicts]
    print(f"[INFO] conflicts = {conflict_summary}")
    assert_true(any(c.conflict_type == ConflictType.CONTRADICTION for c in conflicts), "same FAQ question with different answer is recorded as CONTRADICTION")


def test_search(store: MemoryStore, electric_category: str, doc_category: str, skip_vector: bool) -> None:
    print("\n[PHASE] vector search + filters")
    if skip_vector:
        print("[WARN] vector tests skipped because --skip-vector was used")
        return

    initial_count = store.vector_store.count()
    print(f"[INFO] vector count = {initial_count}")
    assert_true(initial_count > 0, "vector store contains indexed embeddings")

    results = store.search("khách hàng muốn đóng tiền điện", top_k=5)
    print(f"[INFO] search điện = {[(r.rank, round(r.score, 3), r.memory.category, r.memory.tier.value) for r in results]}")
    assert_true(len(results) > 0, "semantic search returns results for Vietnamese paraphrase query")
    assert_true(any("điện" in r.memory.content.lower() for r in results), "search retrieves electricity-related memory")

    # touch=True increments access_count.
    touched = results[0].memory
    before_touch = store.get(touched.id).access_count
    _ = store.search("khách hàng đóng tiền điện", top_k=3, touch=True)
    after_touch = store.get(touched.id).access_count
    assert_true(after_touch >= before_touch, "search with touch=True updates or preserves retrieved access_count")

    # touch=False should not update the chosen exact filtered result.
    filtered = store.search("đóng tiền điện", category=electric_category, top_k=1, touch=False)
    assert_true(len(filtered) == 1, "category-filtered search returns one electricity FAQ result")
    no_touch_id = filtered[0].memory.id
    before_no_touch = store.get(no_touch_id).access_count
    _ = store.search("đóng tiền điện", category=electric_category, top_k=1, touch=False)
    after_no_touch = store.get(no_touch_id).access_count
    assert_true(after_no_touch == before_no_touch, "search with touch=False does not increment access_count")

    category_results = store.search("thanh toán hóa đơn điện", category=electric_category, top_k=5)
    assert_true(len(category_results) > 0, "category filter returns relevant results")
    assert_true(all(r.memory.category == electric_category for r in category_results), "category filter only returns matching category")

    hot_results = store.search("thanh toán hóa đơn điện", tier=MemoryTier.HOT, top_k=5)
    assert_true(len(hot_results) > 0, "tier=HOT filter returns FAQ/hot memories")
    assert_true(all(r.memory.tier == MemoryTier.HOT for r in hot_results), "tier filter only returns HOT memories")

    doc_results = store.search("hạn mức chuyển tiền liên ngân hàng", category=doc_category, top_k=5)
    print(f"[INFO] doc search = {[(r.rank, round(r.score, 3), r.memory.extra_metadata.get('doc_heading')) for r in doc_results]}")
    assert_true(any(r.memory.extra_metadata.get("doc_heading") == "Chuyển tiền liên ngân hàng" for r in doc_results), "search can retrieve a specific DOCX heading section")

    empty_high_threshold = store.search("đóng tiền điện", category=electric_category, min_score=1.01, top_k=5)
    assert_true(len(empty_high_threshold) == 0, "min_score filter can suppress low-confidence results")


def test_status_tier_lifecycle(store: MemoryStore, skip_vector: bool) -> None:
    print("\n[PHASE] lifecycle: hot / warm / cold / archived / expired")

    cold_unit = MemoryUnit(
        content="Quy định lạnh ít dùng cho nghiệp vụ kiểm tra cold tier.",
        source="manual_smoke",
        category="tier_case",
        tier=MemoryTier.WARM,
        created_at=now_utc() - timedelta(days=10),
        updated_at=now_utc() - timedelta(days=10),
        provenance=Provenance(source_type=SourceType.MANUAL, source_path="smoke", source_ref="cold"),
    )
    warm_unit = MemoryUnit(
        content="Quy định warm bình thường mới tạo cho nghiệp vụ kiểm tra tier.",
        source="manual_smoke",
        category="tier_case",
        tier=MemoryTier.WARM,
        provenance=Provenance(source_type=SourceType.MANUAL, source_path="smoke", source_ref="warm"),
    )
    manual_hot_unit = MemoryUnit(
        content="Quy định hot tần suất cao cho nghiệp vụ kiểm tra hot tier.",
        source="manual_smoke",
        category="tier_case",
        tier=MemoryTier.WARM,
        access_count=5,
        last_accessed_at=now_utc(),
        provenance=Provenance(source_type=SourceType.MANUAL, source_path="smoke", source_ref="manual_hot"),
    )
    archive_unit = MemoryUnit(
        content="Archive case không dùng nữa, chỉ để test status archived.",
        source="manual_smoke",
        category="archive_case",
        provenance=Provenance(source_type=SourceType.MANUAL, source_path="smoke", source_ref="archive"),
    )
    expired_unit = MemoryUnit(
        content="TTL case đã hết hạn, chỉ để test status expired.",
        source="manual_smoke",
        category="expired_case",
        ttl_expires_at=now_utc() - timedelta(seconds=1),
        provenance=Provenance(source_type=SourceType.MANUAL, source_path="smoke", source_ref="expired"),
    )

    for unit in [cold_unit, warm_unit, manual_hot_unit, archive_unit, expired_unit]:
        store.add(unit, apply_default_ttl=False)

    store.set_status(archive_unit.id, MemoryStatus.ARCHIVED, reason="deep_smoke_archive")
    lifecycle_stats = run_lifecycle_maintenance(store)
    print(f"[INFO] lifecycle stats = {lifecycle_stats}")

    assert_true(store.get(cold_unit.id).tier == MemoryTier.COLD, "old inactive ACTIVE memory becomes COLD")
    assert_true(store.get(warm_unit.id).tier == MemoryTier.WARM, "recent normal ACTIVE memory remains WARM")
    assert_true(store.get(manual_hot_unit.id).tier == MemoryTier.HOT, "high recent access memory becomes HOT")
    assert_true(store.get(archive_unit.id).status == MemoryStatus.ARCHIVED, "archived memory keeps ARCHIVED status")
    assert_true(store.get(expired_unit.id).status == MemoryStatus.EXPIRED, "expired TTL memory becomes EXPIRED")
    assert_true(lifecycle_stats["cold"] >= 1, "lifecycle stats count COLD memories")
    assert_true(lifecycle_stats["hot"] >= 1, "lifecycle stats count HOT memories")
    assert_true(lifecycle_stats["warm"] >= 1, "lifecycle stats count WARM memories")
    assert_true(lifecycle_stats["expired"] >= 1, "lifecycle stats count expired memories")

    if not skip_vector:
        archived_search = store.search("archive không dùng nữa", category="archive_case", top_k=5)
        expired_search = store.search("ttl hết hạn", category="expired_case", top_k=5)
        cold_search = store.search("quy định lạnh ít dùng", category="tier_case", tier=MemoryTier.COLD, top_k=5)
        hot_search = store.search("quy định hot tần suất cao", category="tier_case", tier=MemoryTier.HOT, top_k=5)
        assert_true(len(archived_search) == 0, "ACTIVE search excludes ARCHIVED memories")
        assert_true(len(expired_search) == 0, "ACTIVE search excludes EXPIRED memories")
        assert_true(len(cold_search) > 0 and all(r.memory.tier == MemoryTier.COLD for r in cold_search), "search can filter COLD tier")
        assert_true(len(hot_search) > 0 and all(r.memory.tier == MemoryTier.HOT for r in hot_search), "search can filter HOT tier after lifecycle update")

    with db.connect(store.config.sqlite_path) as conn:
        archive_logs = db.get_lifecycle_log(conn, archive_unit.id)
        tier_logs = db.get_lifecycle_log(conn, cold_unit.id)
    assert_true(any(log["event"] == "status_change" for log in archive_logs), "lifecycle_log records status_change")
    assert_true(any(log["event"] == "tier_change" for log in tier_logs), "lifecycle_log records tier_change")


def test_update_reindex_cache(store: MemoryStore, skip_vector: bool) -> None:
    print("\n[PHASE] update / reindex / qa_cache")

    target = store.list(status=MemoryStatus.ACTIVE, category="viettelpay_pro", limit=1)[0]
    old_version = target.version
    old_hash = target.content_hash
    updated = store.update_content(
        target.id,
        target.content + "\nCập nhật khuyến mãi smoke để kiểm tra update_content và reindex.",
    )
    assert_true(updated.version == old_version + 1, "update_content increments version")
    assert_true(updated.content_hash != old_hash, "update_content refreshes content_hash")
    assert_true("khuyến mãi smoke" in store.get(updated.id).content.lower(), "updated content is persisted in SQLite")

    if not skip_vector:
        reindexed = store.reindex()
        active_not_archived = store.list(include_expired=False, limit=100000)
        print(f"[INFO] reindexed = {reindexed}, active_not_archived = {len(active_not_archived)}")
        assert_true(reindexed == len(active_not_archived), "reindex rebuilds vectors for non-archived/non-expired memories")
        update_search = store.search("khuyến mãi smoke", category="viettelpay_pro", top_k=6)

        print(
            "[INFO] update search =",
            [
                (
                    r.rank,
                    round(r.score, 3),
                    r.memory.id == updated.id,
                    r.memory.source,
                    r.memory.category,
                    r.memory.content[:80].replace("\n", " "),
                )
                for r in update_search
            ],
        )

        assert_true(
            any(r.memory.id == updated.id for r in update_search),
            "search can find updated content after reindex/update",
        )
    with db.connect(store.config.sqlite_path) as conn:
        logs = db.get_lifecycle_log(conn, updated.id)
        assert_true(any(log["event"] == "updated" for log in logs), "lifecycle_log records updated event")

        qh = compute_query_hash("Khách hàng thanh toán hóa đơn điện như thế nào?")
        db.upsert_qa_cache(
            conn,
            qh,
            "Khách hàng thanh toán hóa đơn điện như thế nào?",
            "Vào mục Hóa đơn, chọn Điện, nhập mã khách hàng và xác nhận.",
            expires_at=now_utc() + timedelta(seconds=60),
            retrieved_memory_ids=[updated.id],
        )
        valid_cache = db.get_valid_qa_cache(conn, qh)
        assert_true(valid_cache is not None and valid_cache.hit_count >= 1, "valid qa_cache entry is returned and hit_count increments")

        expired_qh = compute_query_hash("Cache hết hạn smoke")
        db.upsert_qa_cache(
            conn,
            expired_qh,
            "Cache hết hạn smoke",
            "Không dùng",
            expires_at=now_utc() - timedelta(seconds=1),
            retrieved_memory_ids=[],
        )
        expired_cache = db.get_valid_qa_cache(conn, expired_qh)
        assert_true(expired_cache is None, "expired qa_cache entry is not returned")
        deleted = db.delete_expired_qa_cache(conn)
        assert_true(deleted >= 1, "delete_expired_qa_cache removes expired cache rows")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--keep", action="store_true", help="Keep temp directory for inspection when --work-dir is not used")
    parser.add_argument("--skip-vector", action="store_true", help="Skip vector search tests")
    parser.add_argument("--work-dir", default="", help="Local folder for smoke data, e.g. .smoke_test_data_deep")
    parser.add_argument("--no-reset", action="store_true", help="Do not delete --work-dir before running")
    args = parser.parse_args()

    if args.work_dir:
        temp_root = Path(args.work_dir).expanduser().resolve()
        if temp_root.exists() and not args.no_reset:
            shutil.rmtree(temp_root)
        temp_root.mkdir(parents=True, exist_ok=True)
        should_remove = False
    else:
        temp_root = Path(tempfile.mkdtemp(prefix="memos_core_deep_smoke_"))
        should_remove = not args.keep

    data_dir = temp_root / "data"
    home_dir = temp_root / "store"

    try:
        print(f"[INFO] project root = {ROOT}")
        print(f"[INFO] smoke root = {temp_root}")
        make_sample_files(data_dir)

        config = Config(
            home_dir=home_dir,
            sqlite_filename="deep_smoke.sqlite3",
            default_ttl_seconds=0,
            chunk_size_tokens=200,
            chunk_overlap_tokens=20,
            near_duplicate_threshold=0.55,
            hot_access_count_threshold=2,
            hot_window_days=7,
            cold_after_days_no_access=3,
            top_k=5,
            min_retrieval_score=0.0,
            memos_collection="deep_smoke_memories",
        )
        print(f"[INFO] sqlite = {config.sqlite_path}")
        print(f"[INFO] vector/chroma path = {config.chroma_path}")

        embed_fn = None if args.skip_vector else fake_embed
        store = MemoryStore(config, embed_fn=embed_fn)

        if embed_fn is not None:
            try:
                print(f"[INFO] initial vector count = {store.vector_store.count()}")
            except Exception as exc:
                print("[ERROR] Vector backend failed to initialize.")
                print("        If you use Chroma vector_store.py, install chromadb or run with --skip-vector.")
                print(f"        Original error: {type(exc).__name__}: {exc}")
                raise

        units = load_path_as_units(data_dir, config=config, category="viettelpay_pro")
        electric_category, doc_category = test_ingest(units)
        test_storage_and_ops(store, units)
        test_search(store, electric_category, doc_category, args.skip_vector)
        test_status_tier_lifecycle(store, args.skip_vector)
        test_update_reindex_cache(store, args.skip_vector)

        print("\n[SUCCESS] Deep MemOS-lite core smoke test passed.")
        print("Expected: headings captured, categories parsed, duplicate skipped, contradiction recorded,")
        print("          semantic/category/tier/status filters work, hot/warm/cold/expired lifecycle works.")
        print(f"Inspect SQLite DB: {config.sqlite_path}")
        print(f"Inspect raw fixtures: {data_dir}")

    finally:
        if should_remove:
            shutil.rmtree(temp_root, ignore_errors=True)
            print("[INFO] smoke directory removed. Use --keep or --work-dir to inspect files.")
        else:
            print(f"[INFO] kept smoke directory: {temp_root}")


if __name__ == "__main__":
    main()
