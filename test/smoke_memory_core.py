"""Smoke test for MemOS-lite core modules.

Tests 7 core files:
- config.py
- schemas.py
- db.py
- vector_store.py
- memory_store.py
- ingest.py
- memory_ops.py

Run from project root:
    python scripts/smoke_memory_core.py

The script creates temporary Vietnamese DOCX/XLSX files, ingests them into
MemoryUnit objects, stores them, tests dedup/conflict/lifecycle/search/update,
and then removes nothing so you can inspect the temp SQLite/Chroma directory.
"""

from __future__ import annotations

import argparse
import math
import shutil
import sys
import tempfile
import unicodedata
from datetime import timedelta
from pathlib import Path

# Allow running from scripts/ or project root.
ROOT = Path(__file__).resolve().parents[1] if Path(__file__).resolve().parent.name == "scripts" else Path.cwd()
sys.path.insert(0, str(ROOT))

from src import db
from src.config import Config
from src.ingest import clean_text, load_path_as_units
from src.memory_ops import add_many, run_lifecycle_maintenance
from src.memory_store import MemoryStore
from src.schemas import (
    MemoryStatus,
    MemoryTier,
    MemoryUnit,
    Provenance,
    SourceType,
    compute_query_hash,
    now_utc,
)


def _strip_accents(text: str) -> str:
    text = unicodedata.normalize("NFD", text.lower())
    return "".join(ch for ch in text if unicodedata.category(ch) != "Mn")


def fake_embed(text: str) -> list[float]:
    """Deterministic tiny embedding for smoke tests; no Ollama/model needed.

    It is not a real embedding model. It only gives similar vectors to Vietnamese
    business phrases such as "đóng tiền điện" and "thanh toán hóa đơn điện".
    """

    t = _strip_accents(text)
    features = [
        ("dien" in t or "evn" in t),
        ("hoa don" in t),
        ("thanh toan" in t or "dong tien" in t or "tra tien" in t or "xac nhan" in t),
        ("chuyen tien" in t or "lien ngan hang" in t),
        ("han muc" in t),
        ("phi" in t),
        ("viettelpay" in t or "kpp" in t),
        ("khong the" in t or "loi" in t),
    ]
    vec = [1.0 if flag else 0.0 for flag in features]
    # Add stable character buckets so unrelated text is not all-zero.
    buckets = [0.0] * 8
    for ch in t:
        if ch.isalnum():
            buckets[ord(ch) % len(buckets)] += 1.0
    vec.extend(buckets)
    norm = math.sqrt(sum(x * x for x in vec)) or 1.0
    return [x / norm for x in vec]


def make_sample_files(data_dir: Path) -> None:
    from docx import Document
    from openpyxl import Workbook

    data_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Thanh toán hóa đơn", level=1)
    doc.add_paragraph(
        "Khách hàng có thể thanh toán hóa đơn điện trên ViettelPay Pro bằng cách vào mục Hóa đơn, "
        "chọn Điện, nhập mã khách hàng và xác nhận thanh toán."
    )
    doc.add_heading("Chuyển tiền liên ngân hàng", level=1)
    doc.add_paragraph(
        "Đại lý có thể thực hiện chuyển tiền liên ngân hàng theo hạn mức được cấu hình cho từng kênh bán."
    )
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Loại giao dịch"
    table.cell(0, 1).text = "Ghi chú"
    table.cell(1, 0).text = "Hóa đơn điện"
    table.cell(1, 1).text = "Cần mã khách hàng"
    doc.save(data_dir / "huong_dan_nghiep_vu.docx")

    wb = Workbook()
    ws = wb.active
    ws.title = "FAQ"
    ws.append(["Chủ đề", "Câu hỏi", "Trả lời", "Cấp 1"])
    ws.append([
        "Hóa đơn",
        "Khách hàng thanh toán hóa đơn điện như thế nào?",
        "Vào mục Hóa đơn, chọn Điện, nhập mã khách hàng và xác nhận thanh toán.",
        "Điện",
    ])
    # Exact duplicate: should be skipped.
    ws.append([
        "Hóa đơn",
        "Khách hàng thanh toán hóa đơn điện như thế nào?",
        "Vào mục Hóa đơn, chọn Điện, nhập mã khách hàng và xác nhận thanh toán.",
        "Điện",
    ])
    # Same question, different answer: should be recorded as contradiction/near duplicate.
    ws.append([
        "Hóa đơn",
        "Khách hàng thanh toán hóa đơn điện như thế nào?",
        "Khách hàng không thể thanh toán hóa đơn điện trong hệ thống.",
        "Điện",
    ])
    ws.append([
        "Chuyển tiền",
        "Hạn mức chuyển tiền liên ngân hàng là bao nhiêu?",
        "Hạn mức chuyển tiền phụ thuộc cấu hình của từng kênh bán.",
        "Liên ngân hàng",
    ])
    wb.save(data_dir / "faq_tieng_viet.xlsx")


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)
    print(f"[OK] {message}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--keep", action="store_true", help="Keep temp directory for inspection")
    parser.add_argument("--skip-vector", action="store_true", help="Test SQLite/storage only; skip vector search")
    args = parser.parse_args()

    temp_root = Path(".smoke_test_data").resolve()
    temp_root.mkdir(parents=True, exist_ok=True)
    data_dir = temp_root / "data"
    home_dir = temp_root / "store"

    try:
        print(f"[INFO] temp_root = {temp_root}")
        make_sample_files(data_dir)

        config = Config(
            home_dir=home_dir,
            sqlite_filename="smoke.sqlite3",
            default_ttl_seconds=0,       # avoid accidental expiry during smoke test
            near_duplicate_threshold=0.70,
            hot_access_count_threshold=1,
            cold_after_days_no_access=30,
            top_k=5,
        )
        print(f"[INFO] sqlite = {config.sqlite_path}")
        print(f"[INFO] chroma = {config.chroma_path}")

        embed_fn = None if args.skip_vector else fake_embed
        store = MemoryStore(config, embed_fn=embed_fn)

        if embed_fn is not None:
            try:
                print(f"[INFO] initial vector count = {store.vector_store.count()}")
            except Exception as exc:
                print("[ERROR] Vector backend failed to initialize.")
                print("        If you use current Chroma vector_store.py, install chromadb or run with --skip-vector.")
                print(f"        Original error: {type(exc).__name__}: {exc}")
                raise

        # 1) Ingest Vietnamese DOCX/XLSX.
        units = load_path_as_units(data_dir, config=config, category="viettelpay_pro")
        print(f"[INFO] loaded units = {len(units)}")
        assert_true(len(units) >= 5, "ingest creates MemoryUnit objects from DOCX + XLSX")
        assert_true(any("hóa đơn" in u.content.lower() or "hoá đơn" in u.content.lower() for u in units), "Vietnamese accents are preserved in content")
        assert_true(any(u.provenance and u.provenance.source_type == SourceType.FAQ for u in units), "FAQ provenance is attached")
        assert_true(any(u.provenance and u.provenance.source_type == SourceType.DOC for u in units), "DOC provenance is attached")
        assert_true(clean_text("  Hoá   đơn\u00a0điện  ") == "Hoá đơn điện", "clean_text normalizes whitespace but keeps Vietnamese accents")

        # 2) Add to storage with dedup/conflict.
        stats = add_many(store, units)
        print(f"[INFO] add_many stats = {stats}")
        assert_true(stats["inserted"] >= 4, "memory_store inserts non-duplicate units")
        assert_true(stats["exact_duplicate_skipped"] >= 1, "exact duplicate is skipped by content_hash")
        assert_true(stats["conflicts"] >= 1, "near duplicate / contradiction conflict is recorded")

        all_units = store.list(limit=1000)
        assert_true(len(all_units) == stats["inserted"], "SQLite list_memory_units returns inserted units")
        assert_true(all(u.content_hash for u in all_units), "content_hash exists on stored units")

        conflicts = store.list_conflicts(limit=100)
        print(f"[INFO] conflicts = {[(c.conflict_type.value, c.resolution.value) for c in conflicts]}")
        assert_true(len(conflicts) >= 1, "conflict table can be read")

        # 3) Search/vector path.
        if embed_fn is not None:
            results = store.search("đóng tiền điện ở đâu", top_k=3)
            print(f"[INFO] search results = {[(r.rank, round(r.score, 3), r.memory.source) for r in results]}")
            assert_true(len(results) > 0, "semantic search returns at least one result")
            assert_true(results[0].memory.access_count >= 1, "search touch_access increments access_count")
            assert_true(store.vector_store.count() >= stats["inserted"], "vector_store count is populated")
        else:
            results = []
            print("[WARN] vector search skipped because --skip-vector was used")

        # 4) Update content/version/hash.
        target = results[0].memory if results else all_units[0]
        old_version = target.version
        old_hash = target.content_hash
        updated = store.update_content(target.id, target.content + "\nCập nhật smoke test.")
        assert_true(updated.version == old_version + 1, "update_content increments version")
        assert_true(updated.content_hash != old_hash, "update_content refreshes content_hash")

        # 5) Status/tier/lifecycle.
        store.set_tier(updated.id, MemoryTier.HOT, reason="smoke_test_hot")
        assert_true(store.get(updated.id).tier == MemoryTier.HOT, "set_tier updates serving tier")

        archive_target = next(u for u in all_units if u.id != updated.id)
        store.set_status(archive_target.id, MemoryStatus.ARCHIVED, reason="smoke_test_archive")
        assert_true(store.get(archive_target.id).status == MemoryStatus.ARCHIVED, "set_status archives a unit")

        expired_unit = MemoryUnit(
            content="Tri thức test TTL đã hết hạn.",
            source="manual_smoke",
            category="ttl_test",
            provenance=Provenance(source_type=SourceType.MANUAL, source_path="smoke", source_ref="ttl"),
            ttl_expires_at=now_utc() - timedelta(seconds=1),
        )
        store.add(expired_unit, apply_default_ttl=False)
        lifecycle_stats = run_lifecycle_maintenance(store)
        print(f"[INFO] lifecycle stats = {lifecycle_stats}")
        assert_true(lifecycle_stats["expired"] >= 1, "expired TTL units are marked EXPIRED")
        assert_true(store.get(expired_unit.id).status == MemoryStatus.EXPIRED, "expired unit status is EXPIRED")

        # 6) Low-level DB helpers: lifecycle log + QA cache.
        with db.connect(config.sqlite_path) as conn:
            logs = db.get_lifecycle_log(conn, updated.id)
            qh = compute_query_hash("Khách hàng thanh toán hóa đơn điện như thế nào?")
            db.upsert_qa_cache(
                conn,
                qh,
                "Khách hàng thanh toán hóa đơn điện như thế nào?",
                "Vào mục Hóa đơn, chọn Điện, nhập mã khách hàng và xác nhận.",
                expires_at=now_utc() + timedelta(seconds=60),
                retrieved_memory_ids=[updated.id],
            )
            cache = db.get_valid_qa_cache(conn, qh)
        assert_true(len(logs) >= 1, "lifecycle_log is readable")
        assert_true(cache is not None and cache.hit_count >= 1, "qa_cache insert/get_valid increments hit_count")

        print("\n[SUCCESS] MemOS-lite core smoke test passed.")
        print("Expected high-level output: loaded units > 0, duplicate skipped >= 1, conflicts >= 1, lifecycle expired >= 1.")
        print(f"Inspect DB: {config.sqlite_path}")

    finally:
        if args.keep:
            print(f"[INFO] kept temp directory: {temp_root}")
        else:
            shutil.rmtree(temp_root, ignore_errors=True)
            print("[INFO] temp directory removed. Use --keep to inspect DB/Chroma files.")


if __name__ == "__main__":
    main()
